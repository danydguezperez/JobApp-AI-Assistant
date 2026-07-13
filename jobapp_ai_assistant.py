from __future__ import annotations

import json
import os
import re
import shutil
import sqlite3
import sys
import tempfile
import textwrap
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional
from xml.sax.saxutils import escape
from xml.etree import ElementTree

import httpx
import requests
from bs4 import BeautifulSoup
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # pragma: no cover - optional dependency path
    Document = None
    Pt = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency path
    PdfReader = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except Exception:  # pragma: no cover - optional dependency path
    A4 = None
    getSampleStyleSheet = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None


if getattr(sys, "frozen", False):
    _EXE_DIR = Path(sys.executable).parent
    BASE_DIR = Path(getattr(sys, "_MEIPASS"))
else:
    _EXE_DIR = Path(__file__).resolve().parent
    BASE_DIR = _EXE_DIR

STATIC_DIR = BASE_DIR / "static"
DATA_DIR = _EXE_DIR / "data"
EXPORTS_DIR = _EXE_DIR / "exports"
DB_PATH = _EXE_DIR / "applications.db"
CV_JSON_PATH = DATA_DIR / "cv_profile.json"
LLM_CONFIG_PATH = DATA_DIR / "llm_providers.json"

def configured_path(env_name: str, fallback: Path) -> Path:
    value = os.getenv(env_name, "").strip()
    return Path(value) if value else fallback


LOCAL_CV_TEXT = configured_path("JOBAPP_LOCAL_CV_TEXT", _EXE_DIR / "data" / "CV_text.txt")
LOCAL_CV_PDF = configured_path("JOBAPP_LOCAL_CV_PDF", _EXE_DIR / "data" / "CV.pdf")
LOCAL_CV_XML = configured_path("JOBAPP_LOCAL_CV_XML", _EXE_DIR / "data" / "CV.xml")
LOCAL_CV_PDF_CANDIDATES = [
    LOCAL_CV_PDF,
    _EXE_DIR / "data" / "CV_DDP_2026.pdf",
    _EXE_DIR / "CV_DDP_2026.pdf",
]
LOCAL_CV_TEXT_CANDIDATES = [
    LOCAL_CV_TEXT,
    _EXE_DIR / "data" / "CV_DDP_2026_text.txt",
    _EXE_DIR / "CV_DDP_2026_text.txt",
]
GEMINI_CONFIG = configured_path("JOBAPP_GEMINI_CONFIG", _EXE_DIR / "config.json")
PERSONA_MASTER = configured_path("JOBAPP_PERSONA_MASTER", _EXE_DIR / "data" / "persona_master_EN.md")
PROFILE_DOSSIER = configured_path("JOBAPP_PROFILE_DOSSIER", _EXE_DIR / "data" / "profile_brief.md")

APP_NAME = "JobApp AI Assistant"
GEMINI_DEFAULT_MODEL = "gemini-3.5-flash"
OPENAI_DEFAULT_MODEL = "gpt-5.6-terra"
ANTHROPIC_DEFAULT_MODEL = "claude-sonnet-5"


app = FastAPI(
    title=APP_NAME,
    description="Local-first CV parsing, job matching, and application drafting assistant.",
    version="0.2.0",
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8080",
        "http://127.0.0.1:8080",
        "http://localhost:8090",
        "http://127.0.0.1:8090",
        "http://localhost:8091",
        "http://127.0.0.1:8091",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class MatchRequest(BaseModel):
    job_text: str = Field(..., min_length=40)
    role_type: str = Field(default="Hybrid")
    selected_cv: dict[str, Any] = Field(default_factory=dict)
    job_url: Optional[str] = None
    company: Optional[str] = None
    position: Optional[str] = None


class ExportRequest(BaseModel):
    adapted_cv: str = ""
    cover_letter: str = ""
    interview_questions: list[Any] | str = Field(default_factory=list)
    position: Optional[str] = None
    company: Optional[str] = None
    format: str = "docx"


class CVExportRequest(BaseModel):
    cv: dict[str, Any] = Field(default_factory=dict)
    format: str = "md"
    filename_prefix: str = "parsed_cv"


class ProviderSaveRequest(BaseModel):
    active_provider: str = "gemini"
    provider: str = "gemini"
    model: str = ""
    api_key: Optional[str] = None
    base_url: Optional[str] = None
    clear_api_key: bool = False


class ProviderTestRequest(BaseModel):
    provider: Optional[str] = None


def ensure_dirs() -> None:
    if getattr(sys, "frozen", False) and not DATA_DIR.exists():
        bundled_data = BASE_DIR / "data"
        if bundled_data.exists():
            shutil.copytree(bundled_data, DATA_DIR, dirs_exist_ok=True)
    DATA_DIR.mkdir(exist_ok=True)
    EXPORTS_DIR.mkdir(exist_ok=True)


def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    ensure_dirs()
    with db() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS applications (
                id TEXT PRIMARY KEY,
                created_at TEXT NOT NULL,
                company TEXT,
                position TEXT,
                role_type TEXT,
                job_url TEXT,
                job_text TEXT NOT NULL,
                result_json TEXT NOT NULL
            )
            """
        )
        conn.commit()


def read_text_file(path: Path, limit: Optional[int] = None) -> str:
    if not path.exists():
        return ""
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = repair_mojibake(path.read_text(encoding=encoding, errors="replace"))
            return text[:limit] if limit else text
        except Exception:
            continue
    return ""


def first_existing_path(paths: list[Path]) -> Optional[Path]:
    for path in paths:
        if path.exists():
            return path
    return None


def load_local_cv_text() -> tuple[str, str]:
    pdf_path = first_existing_path(LOCAL_CV_PDF_CANDIDATES)
    if pdf_path:
        return extract_pdf_path(pdf_path), str(pdf_path)
    text_path = first_existing_path(LOCAL_CV_TEXT_CANDIDATES)
    if text_path:
        return read_text_file(text_path), str(text_path)
    return "", ""


def repair_mojibake(text: str) -> str:
    markers = ("Ã", "Â", "â€", "â€“", "â€¢", "ðŸ")
    if not any(marker in text for marker in markers):
        return text
    try:
        repaired = text.encode("latin-1", errors="ignore").decode("utf-8", errors="ignore")
    except Exception:
        return text
    old_score = sum(text.count(marker) for marker in markers)
    new_score = sum(repaired.count(marker) for marker in markers)
    return repaired if new_score < old_score else text


def repair_data(value: Any) -> Any:
    if isinstance(value, str):
        return repair_mojibake(value)
    if isinstance(value, list):
        return [repair_data(item) for item in value]
    if isinstance(value, dict):
        return {key: repair_data(item) for key, item in value.items()}
    return value


def normalize_cv_profile(cv: dict[str, Any]) -> dict[str, Any]:
    cv = repair_data(cv)
    personal = cv.get("personal_info")
    if isinstance(personal, dict):
        citation_names = personal.get("citation_names")
        if isinstance(citation_names, str):
            citation_names = re.split(r"\s*[,;\n]\s*|(?<=\b[A-Z])(?=Dany\s)", citation_names)
        if isinstance(citation_names, list):
            cleaned = []
            for value in citation_names:
                name = re.sub(r"\s+", " ", str(value)).strip(" ,;")
                if not name:
                    continue
                if re.search(r"\bD{2}any\b", name, flags=re.IGNORECASE):
                    name = re.sub(r"\bD(?=Dany\b)", "", name)
                for part in re.split(r"(?<=\b[A-Z])(?=Dany\s)", name):
                    part = part.strip(" ,;")
                    if part and part not in cleaned:
                        cleaned.append(part)
            if cleaned:
                personal["citation_names"] = cleaned
    return cv


def extract_xml_text(path: Path, limit: int = 20000) -> str:
    if not path.exists():
        return ""
    try:
        root = ElementTree.parse(path).getroot()
        parts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                value = re.sub(r"\s+", " ", elem.text.strip())
                if value not in parts:
                    parts.append(value)
            if len(" ".join(parts)) > limit:
                break
        return "\n".join(parts)[:limit]
    except Exception:
        return read_text_file(path, limit=limit)


def extract_xml_text_from_string(raw: str, limit: int = 80000) -> str:
    try:
        root = ElementTree.fromstring(raw)
        parts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                value = re.sub(r"\s+", " ", elem.text.strip())
                if value not in parts:
                    parts.append(value)
            if len(" ".join(parts)) > limit:
                break
        return "\n".join(parts)[:limit]
    except Exception:
        return raw[:limit]


def load_gemini_config() -> tuple[Optional[str], str]:
    config = {}
    if GEMINI_CONFIG.exists():
        try:
            config = json.loads(read_text_file(GEMINI_CONFIG))
        except Exception:
            config = {}
    api_key = (
        os.getenv("GEMINI_API_KEY")
        or config.get("gemini_api_key")
        or config.get("api_key")
        or config.get("GOOGLE_API_KEY")
    )
    model = config.get("preferred_model") or os.getenv("GEMINI_MODEL") or GEMINI_DEFAULT_MODEL
    return api_key, model


def default_llm_settings() -> dict[str, Any]:
    api_key, model = load_gemini_config()
    return {
        "active_provider": "gemini",
        "providers": {
            "gemini": {
                "label": "Google Gemini",
                "api_key": api_key or "",
                "model": model or GEMINI_DEFAULT_MODEL,
                "base_url": "https://generativelanguage.googleapis.com/v1beta",
            },
            "openai": {
                "label": "OpenAI / ChatGPT",
                "api_key": os.getenv("OPENAI_API_KEY", ""),
                "model": OPENAI_DEFAULT_MODEL,
                "base_url": "https://api.openai.com/v1",
            },
            "anthropic": {
                "label": "Anthropic Claude",
                "api_key": os.getenv("ANTHROPIC_API_KEY", ""),
                "model": ANTHROPIC_DEFAULT_MODEL,
                "base_url": "https://api.anthropic.com/v1",
            },
            "openai_compatible": {
                "label": "Custom OpenAI-compatible API",
                "api_key": "",
                "model": "model-name",
                "base_url": "http://localhost:11434/v1",
            },
            "local_openai": {
                "label": "Local OpenAI-compatible server",
                "api_key": "",
                "model": "local-model",
                "base_url": "http://localhost:8081/v1",
            },
            "ollama": {
                "label": "Ollama local",
                "api_key": "",
                "model": "llama3.1",
                "base_url": "http://localhost:11434/v1",
            },
            "deepseek": {
                "label": "DeepSeek",
                "api_key": os.getenv("DEEPSEEK_API_KEY", ""),
                "model": "deepseek-chat",
                "base_url": "https://api.deepseek.com/v1",
            },
            "glm": {
                "label": "GLM / Zhipu AI",
                "api_key": os.getenv("GLM_API_KEY", ""),
                "model": "glm-4.5",
                "base_url": "https://open.bigmodel.cn/api/paas/v4",
            },
        },
    }


def load_llm_settings() -> dict[str, Any]:
    defaults = default_llm_settings()
    if not LLM_CONFIG_PATH.exists():
        return defaults
    try:
        saved = json.loads(read_text_file(LLM_CONFIG_PATH))
    except Exception:
        return defaults
    merged = defaults
    merged["active_provider"] = saved.get("active_provider") or merged["active_provider"]
    for name, provider in saved.get("providers", {}).items():
        if name not in merged["providers"]:
            merged["providers"][name] = {}
        merged["providers"][name].update(provider or {})
    return merged


def save_llm_settings(settings: dict[str, Any]) -> None:
    ensure_dirs()
    LLM_CONFIG_PATH.write_text(json.dumps(settings, ensure_ascii=False, indent=2), encoding="utf-8")


def public_llm_settings(settings: Optional[dict[str, Any]] = None) -> dict[str, Any]:
    settings = settings or load_llm_settings()
    public = {"active_provider": settings.get("active_provider", "gemini"), "providers": {}}
    for name, provider in settings.get("providers", {}).items():
        key = provider.get("api_key") or ""
        public["providers"][name] = {
            "label": provider.get("label") or name,
            "model": provider.get("model") or "",
            "base_url": provider.get("base_url") or "",
            "api_key_present": bool(key),
            "api_key_mask": f"{key[:4]}...{key[-4:]}" if len(key) > 8 else ("configured" if key else ""),
        }
    return public


def provider_can_run_without_key(provider_name: str) -> bool:
    return provider_name in {"openai_compatible", "local_openai", "ollama"}


def active_provider_config(provider_name: Optional[str] = None) -> tuple[str, dict[str, Any]]:
    settings = load_llm_settings()
    name = provider_name or settings.get("active_provider", "gemini")
    provider = settings.get("providers", {}).get(name)
    if not provider:
        raise RuntimeError(f"LLM provider '{name}' is not configured.")
    return name, provider


def llm_generate(prompt: str, temperature: float = 0.25, provider_name: Optional[str] = None) -> str:
    name, provider = active_provider_config(provider_name)
    if name == "gemini":
        return gemini_generate(prompt, temperature, provider)
    if name == "openai":
        return openai_generate(prompt, temperature, provider, use_responses=True)
    if name == "anthropic":
        return anthropic_generate(prompt, temperature, provider)
    if name in {"openai_compatible", "local_openai", "ollama", "deepseek", "glm"}:
        return openai_generate(prompt, temperature, provider, use_responses=False)
    raise RuntimeError(f"Unsupported LLM provider '{name}'.")


def gemini_generate(prompt: str, temperature: float = 0.25, provider: Optional[dict[str, Any]] = None) -> str:
    provider = provider or active_provider_config("gemini")[1]
    api_key = provider.get("api_key")
    model = provider.get("model") or GEMINI_DEFAULT_MODEL
    base_url = (provider.get("base_url") or "https://generativelanguage.googleapis.com/v1beta").rstrip("/")
    if not api_key:
        raise RuntimeError("Gemini API key was not found. Configure it in Motor IA or config.json.")

    url = f"{base_url}/models/{model}:generateContent"
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": temperature,
            "topP": 0.9,
            "maxOutputTokens": 30000,
            "responseMimeType": "application/json",
        },
    }
    response = requests.post(url, params={"key": api_key}, json=payload, timeout=240)
    if response.status_code >= 400:
        raise RuntimeError(f"Gemini request failed with HTTP {response.status_code}: {response.text[:500]}")
    data = response.json()
    try:
        return "\n".join(
            part.get("text", "")
            for candidate in data.get("candidates", [])
            for part in candidate.get("content", {}).get("parts", [])
            if part.get("text")
        ).strip()
    except Exception as exc:
        raise RuntimeError(f"Unexpected Gemini response format: {exc}") from exc


def openai_generate(
    prompt: str,
    temperature: float = 0.25,
    provider: Optional[dict[str, Any]] = None,
    use_responses: bool = True,
) -> str:
    provider = provider or active_provider_config("openai")[1]
    api_key = provider.get("api_key")
    model = provider.get("model") or OPENAI_DEFAULT_MODEL
    base_url = (provider.get("base_url") or "https://api.openai.com/v1").rstrip("/")
    if not api_key and use_responses:
        raise RuntimeError("OpenAI API key was not found. Configure it in Motor IA or OPENAI_API_KEY.")

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    if use_responses:
        payload = {
            "model": model,
            "input": prompt,
            "temperature": temperature,
            "max_output_tokens": 24000,
        }
        response = requests.post(f"{base_url}/responses", headers=headers, json=payload, timeout=240)
        if response.status_code == 404:
            return openai_chat_generate(prompt, temperature, provider)
        if response.status_code >= 400:
            raise RuntimeError(f"OpenAI request failed with HTTP {response.status_code}: {response.text[:500]}")
        data = response.json()
        if data.get("output_text"):
            return data["output_text"].strip()
        parts = []
        for item in data.get("output", []):
            for content in item.get("content", []):
                if content.get("text"):
                    parts.append(content["text"])
        return "\n".join(parts).strip()
    return openai_chat_generate(prompt, temperature, provider)


def openai_chat_generate(prompt: str, temperature: float, provider: dict[str, Any]) -> str:
    api_key = provider.get("api_key")
    model = provider.get("model") or "model-name"
    base_url = (provider.get("base_url") or "https://api.openai.com/v1").rstrip("/")
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": 24000,
    }
    response = requests.post(f"{base_url}/chat/completions", headers=headers, json=payload, timeout=240)
    if response.status_code >= 400:
        raise RuntimeError(f"OpenAI-compatible request failed with HTTP {response.status_code}: {response.text[:500]}")
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


def anthropic_generate(prompt: str, temperature: float = 0.25, provider: Optional[dict[str, Any]] = None) -> str:
    provider = provider or active_provider_config("anthropic")[1]
    api_key = provider.get("api_key")
    model = provider.get("model") or ANTHROPIC_DEFAULT_MODEL
    base_url = (provider.get("base_url") or "https://api.anthropic.com/v1").rstrip("/")
    if not api_key:
        raise RuntimeError("Anthropic API key was not found. Configure it in Motor IA or ANTHROPIC_API_KEY.")

    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": model,
        "max_tokens": 24000,
        "messages": [{"role": "user", "content": prompt}],
    }
    if not model.startswith("claude-sonnet-5"):
        payload["temperature"] = temperature
    response = requests.post(f"{base_url}/messages", headers=headers, json=payload, timeout=240)
    if response.status_code >= 400:
        raise RuntimeError(f"Anthropic request failed with HTTP {response.status_code}: {response.text[:500]}")
    data = response.json()
    return "\n".join(part.get("text", "") for part in data.get("content", []) if part.get("text")).strip()


def strip_json_fence(text: str) -> str:
    text = text.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1).strip()
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        return text[start : end + 1]
    return text


def parse_json_response(text: str) -> dict[str, Any]:
    return json.loads(strip_json_fence(text))


def llm_json(prompt: str, temperature: float = 0.1) -> dict[str, Any]:
    raw = llm_generate(prompt, temperature=temperature)
    try:
        return parse_json_response(raw)
    except json.JSONDecodeError as exc:
        repair_prompt = f"""
Repair the following invalid JSON into valid JSON.
Rules:
- Return only valid JSON.
- Preserve all fields and content.
- Do not summarize.
- Do not add commentary.

JSON error: {exc}

Invalid JSON:
{raw[:60000]}
"""
        repaired = llm_generate(repair_prompt, temperature=0)
        return parse_json_response(repaired)


def extract_pdf_text(upload: UploadFile) -> str:
    if PdfReader is None:
        raise HTTPException(status_code=500, detail="pypdf is not installed. Run: pip install -r requirements.txt")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(upload.file.read())
        tmp_path = tmp.name
    try:
        reader = PdfReader(tmp_path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages).strip()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def extract_pdf_path(path: Path) -> str:
    if not path.exists():
        return ""
    if PdfReader is None:
        raise HTTPException(status_code=500, detail="pypdf is not installed. Run: pip install -r requirements.txt")
    reader = PdfReader(str(path))
    pages = [f"--- Page {index + 1} ---\n{page.extract_text() or ''}" for index, page in enumerate(reader.pages)]
    return repair_mojibake("\n\n".join(pages).strip())


def chunk_words(text: str, max_words: int = 40) -> str:
    words = re.sub(r"\s+", " ", text).strip().split()
    return " ".join(words[:max_words])


def heuristic_cv_json(cv_text: str, dossier: str = "") -> dict[str, Any]:
    merged = f"{cv_text}\n{dossier}"
    email = re.search(r"[\w.\-+]+@[\w.\-]+\.\w+", merged)
    phone = re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", merged)
    skills = [
        "Bioinformatics", "Proteo-transcriptomics", "RNA-Seq", "Single-cell omics",
        "De novo transcriptome assembly", "Shotgun proteomics", "Proteogenomics",
        "Python", "R", "Bash", "Biopython", "Snakemake", "Nextflow",
        "Proteomics", "Mass spectrometry", "NGS pipeline development",
        "Phylogenetics", "IQ-TREE", "AlphaFold2", "Biodiscovery",
        "Bioactive peptide discovery", "Antimicrobial peptides", "Toxinology",
        "B2B scientific sales", "Account management", "Scientific consulting",
        "Project leadership", "Academic mentoring", "Scientific writing",
        "Peer review", "Systematic review", "Biostatistics",
    ]
    profile = {
        "personal_info": {
            "full_name": "Dany Domínguez Pérez",
            "headline": "PhD Bioinformatician | RNA-seq · Proteomics · NGS · Pipeline Development",
            "email": email.group(0) if email else "danydguezperez@gmail.com",
            "phone": phone.group(0).strip() if phone else "+351 911 976 480",
            "location": "Porto, Portugal",
            "links": [
                "https://orcid.org/0000-0002-5211-972X",
                "https://www.researchgate.net/profile/Dany_Dominguez_Perez3",
                "https://www.linkedin.com/in/dany-dom%C3%ADnguez-p%C3%A9rez-755b6381",
                "https://www.pagbiomics.com",
            ],
            "summary": (
                "PhD bioinformatician with 10+ years of experience transforming raw biological data "
                "into publication-ready results. Lead author of DeTox (Briefings in Bioinformatics, 2024) "
                "and SeqLengthPlot v2.0 (Bioinformatics Advances, 2024). 43 peer-reviewed publications, "
                "413 citations. Specializes in proteotranscriptomics, NGS pipelines, and bioactive peptide "
                "discovery for non-model and marine organisms. Runs private dual-Xeon HPC (160 threads, "
                "384 GB RAM). Available for freelance and consulting. Trilingual: Spanish (native), "
                "Portuguese (fluent), English (professional)."
            ),
        },
        "education": [
            {
                "degree": "PhD in Biology",
                "institution": "University of Porto, Faculty of Sciences",
                "year": "2017",
                "description": (
                    "Thesis on proteotranscriptomics and venom characterization of marine organisms. "
                    "Integrated de novo transcriptome assembly, shotgun proteomics, and phylogenetic analysis."
                ),
            },
            {
                "degree": "Biology Degree (Licenciatura)",
                "institution": "University of Havana, Cuba",
                "year": "2009",
                "description": "Specialization in marine biology and early toxinology.",
            },
            {
                "degree": "Diploma in Higher Education (Pedagogical)",
                "institution": "Universidad Central Marta Abreu de Las Villas, Cuba",
                "year": "2011",
                "description": "University teaching qualification.",
            },
            {
                "degree": "Postgraduate Course — Single-Cell Technology",
                "institution": "The Single-Cell World",
                "year": "2023",
                "description": "Advanced training in single-cell omics methodologies.",
            },
        ],
        "experience": [
            {
                "title": "Independent Bioinformatician & Freelancer",
                "company": "PagBiOMICs / Upwork",
                "years": "2026–present",
                "description": (
                    "Freelance bioinformatics consulting. Services: RNA-seq analysis, proteomics, "
                    "NGS pipeline development, phylogenetics, bioactive peptide discovery, "
                    "scientific manuscript review. Operates dual-Xeon HPC cluster (160 threads, 384 GB RAM)."
                ),
            },
            {
                "title": "Research Fellow / Postdoctoral Researcher",
                "company": "Stazione Zoologica Anton Dohrn (SZN), Naples, Italy",
                "years": "Feb 2024 – Jan 2026",
                "description": (
                    "CRIMAC DEEPVEN project (funded by Italian MUR). Proteotranscriptomic characterization "
                    "of deep-sea Anthozoa venoms. Built DeTox pipeline (Briefings in Bioinformatics, 2024). "
                    "Integrated AlphaFold2 structural prediction, mass spectrometry-guided proteomics, "
                    "and ion-channel pharmacological profiling for toxin discovery."
                ),
            },
            {
                "title": "Account Manager / Scientific Sales Representative",
                "company": "Proquinorte, Unipessoal Lda",
                "years": "2022–2023",
                "description": (
                    "B2B scientific equipment and diagnostics sales in Portugal. Managed laboratory "
                    "accounts across industrial, clinical, and research sectors. Technical translation "
                    "of product value to scientific decision-makers."
                ),
            },
            {
                "title": "Assistant Researcher & Co-Manager",
                "company": "CIIMAR — Interdisciplinary Centre of Marine and Environmental Research, University of Porto",
                "years": "2018–2022",
                "description": (
                    "MOREBIVALVES project. Led RNA-seq differential expression and quantitative proteomics "
                    "of bivalves (cockles, mussels) under marine biotoxin exposure. Co-managed lab operations "
                    "and supervised MSc students."
                ),
            },
            {
                "title": "Research Assistant / Proteomics Platform Coordinator",
                "company": "EMBRC-PT — European Marine Biological Resource Centre, Portugal",
                "years": "2017–2018",
                "description": "Coordinated molecular biology and mass spectrometry platform for marine biological resources.",
            },
            {
                "title": "Lecturer, Department of Biology",
                "company": "Universidad Central Marta Abreu de Las Villas, Cuba",
                "years": "2009–2012",
                "description": "Taught molecular biology, genetics, ecology, and cell biology at undergraduate level.",
            },
        ],
        "publications": [
            {
                "title": "DeTox: a pipeline for the detection of toxins in venomous organisms",
                "journal": "Briefings in Bioinformatics",
                "year": "2024",
                "role": "Lead author",
            },
            {
                "title": "SeqLengthPlot v2.0: interactive visualization of FASTA sequence length distributions",
                "journal": "Bioinformatics Advances",
                "year": "2024",
                "role": "Lead author",
            },
            {
                "title": "Unveiling Encrypted Antimicrobial Peptides from Cephalopods' Salivary Glands",
                "journal": "ACS Omega",
                "year": "2024",
                "role": "Co-author",
            },
            {
                "title": "43 peer-reviewed publications (413 citations) across genomics, proteomics, toxinology, marine biology",
                "journal": "Multiple indexed journals including Toxins, npj Biofilms & Microbiomes, IJMS, Frontiers in Microbiology",
                "year": "2009–2026",
                "role": "Author/co-author",
            },
        ],
        "projects": [
            {
                "name": "PagBiOMICs",
                "role": "Founder",
                "description": (
                    "Bioinformatics consulting platform and brand manager. Services: OMICs analysis, "
                    "pipeline development, scientific writing. Active distributor sourcing pipeline "
                    "(41 verified brands). FastAPI backend with local dashboard."
                ),
            },
            {
                "name": "CRIMAC DEEPVEN",
                "role": "Postdoctoral lead",
                "description": "Deep-sea Anthozoa venom discovery project at SZN Naples. Produced DeTox pipeline and 3 publications.",
            },
            {
                "name": "MOREBIVALVES",
                "role": "Scientific lead / Co-manager",
                "description": "Molecular toxicology, transcriptomics, and proteomics of commercial bivalves under biotoxin exposure.",
            },
        ],
        "skills": skills,
        "software": [
            "FiltDeTox (author · github.com/danydguezperez/FiltDeTox)",
            "DeTox (lead author, Briefings in Bioinformatics 2024)",
            "SeqLengthPlot v2.0 (lead author)",
            "FastAPI",
            "Snakemake",
            "Nextflow",
            "MaxQuant",
            "Perseus",
            "FragPipe",
            "DESeq2",
            "IQ-TREE",
            "BEAST",
        ],
        "languages": [
            {"language": "Spanish", "level": "Native"},
            {"language": "Portuguese", "level": "Fluent"},
            {"language": "English", "level": "Professional"},
        ],
        "infrastructure": "Dual-node Xeon HPC: 160 threads, 384 GB RAM, 9.3 TB storage, 1 Gbps (Porto, Portugal)",
        "source_notes": {
            "mode": "heuristic_fallback",
            "cv_excerpt": chunk_words(cv_text, 80),
        },
    }
    return profile


def compact_cv_text(text: str) -> str:
    text = re.sub(r"© CIÊNCIAVITAE.*?Page \d+", " ", text)
    text = re.sub(r"--- Page \d+ ---", " ", text)
    text = re.sub(r"\nDany Dom[íi]nguez P[ée]rez\n", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def section_between(text: str, start_pattern: str, end_patterns: list[str]) -> str:
    start = re.search(start_pattern, text, flags=re.IGNORECASE)
    if not start:
        return ""
    tail = text[start.start() :]
    ends = [m.start() for pattern in end_patterns if (m := re.search(pattern, tail, flags=re.IGNORECASE))]
    end = min(ends) if ends else len(tail)
    return tail[:end]


def local_publications(cv_text: str) -> list[dict[str, str]]:
    text = compact_cv_text(cv_text)
    outputs = section_between(text, r"\bOutputs\s+Publications\b", [r"\bActivities\b"])
    if not outputs:
        outputs = section_between(text, r"\bPublications\b", [r"\bActivities\b", r"\bSupervision\b"])
    if not outputs:
        return []

    category_marks = []
    for category in ("Book chapter", "Conference abstract", "Conference poster", "Journal article"):
        for match in re.finditer(re.escape(category), outputs, flags=re.IGNORECASE):
            category_marks.append((match.start(), category))
    category_marks.sort()

    publications = []
    seen = set()
    for match in re.finditer(r'"([^"]{18,700})"', outputs):
        title = re.sub(r"\s+", " ", match.group(1)).strip(" .")
        key = title.lower()
        if key in seen:
            continue
        seen.add(key)
        before = outputs[max(0, match.start() - 450) : match.start()]
        after = outputs[match.end() : match.end() + 700]
        context = f"{before} {title} {after}"
        years = re.findall(r"\b(?:19|20)\d{2}\b", context)
        doi = re.search(r"\b10\.\d{4,9}/[^\s,;]+", context)
        category = "Publication"
        for pos, candidate in category_marks:
            if pos <= match.start():
                category = candidate
            else:
                break
        journal_candidate = re.sub(r"\s+", " ", after.strip()).split(". ")[0][:180]
        publications.append(
            {
                "title": title,
                "journal": journal_candidate or category,
                "year": years[-1] if years else "",
                "type": category,
                "doi": doi.group(0).rstrip(".") if doi else "",
            }
        )
    return publications


def local_activity_rows(cv_text: str, heading: str, limit: int = 12) -> list[dict[str, str]]:
    text = compact_cv_text(cv_text)
    section = section_between(text, rf"\b{re.escape(heading)}\b", [r"\bDistinctions\b", r"\bOutputs\b", r"\bProjects\b"])
    if not section:
        return []
    rows = []
    for match in re.finditer(r"\b(?:19|20)\d{2}(?:/\d{2}/\d{2})?(?:\s*-\s*(?:19|20)\d{2}(?:/\d{2}/\d{2})?)?", section):
        snippet = re.sub(r"\s+", " ", section[match.start() : match.start() + 420]).strip()
        if snippet and snippet not in [row["description"] for row in rows]:
            rows.append({"description": snippet})
        if len(rows) >= limit:
            break
    return rows


def local_supervision_rows(cv_text: str) -> list[dict[str, str]]:
    text = compact_cv_text(cv_text)
    section = section_between(text, r"\bSupervision\s+Thesis Title\b", [r"\bEvent participation\b", r"\bDistinctions\b"])
    rows = []
    for match in re.finditer(r"\b(?:19|20)\d{2}/\d{2}/\d{2}\s*-\s*(?:Current|(?:19|20)\d{2}/\d{2}/\d{2})", section):
        snippet = re.sub(r"\s+", " ", section[match.start() : match.start() + 520]).strip()
        if snippet:
            rows.append({"description": snippet})
    return rows


def enrich_cv_from_source(parsed: dict[str, Any], cv_text: str) -> dict[str, Any]:
    publications = local_publications(cv_text)
    if len(publications) > len(parsed.get("publications") or []):
        parsed["publications"] = publications
    supervision_rows = local_supervision_rows(cv_text)
    if supervision_rows and len(supervision_rows) > len(parsed.get("supervision") or []):
        parsed["supervision"] = supervision_rows
    for field, heading in (
        ("events", "Oral presentation"),
        ("awards", "Distinctions"),
    ):
        local_rows = local_activity_rows(cv_text, heading)
        if local_rows and len(local_rows) > len(parsed.get(field) or []):
            parsed[field] = local_rows
    parsed.setdefault("source_notes", {})
    parsed["source_notes"].update(
        {
            "source_mode": "pdf_primary",
            "source_chars": len(cv_text),
            "local_publications_found": len(publications),
        }
    )
    return parsed


def load_reference_bundle() -> dict[str, str]:
    # Persona master is primary; fall back to Antigravity dossier if it exists
    persona = read_text_file(PERSONA_MASTER)
    dossier = read_text_file(PROFILE_DOSSIER)
    combined_dossier = persona if persona else dossier
    if persona and dossier:
        combined_dossier = f"{persona}\n\n---\n\n{dossier}"
    return {
        "cv_text": read_text_file(LOCAL_CV_TEXT),
        "cv_xml_text": extract_xml_text(LOCAL_CV_XML),
        "dossier": combined_dossier,
    }


def cv_to_structured_json(cv_text: str) -> dict[str, Any]:
    refs = load_reference_bundle()
    dossier = refs["dossier"][:5000]
    xml_text = refs["cv_xml_text"][:12000]
    outputs_match = re.search(r"\bOutputs\s+Publications\b", cv_text, flags=re.IGNORECASE)
    if outputs_match:
        cv_payload = cv_text[: outputs_match.start()]
    else:
        cv_payload = cv_text[:80000]
    prompt = f"""
You are an expert CV parser. Your job is extraction, not summarization.

Task:
Convert the SOURCE CV into strict JSON in English. Preserve factual accuracy and maximum useful detail for ATS matching.
Return only valid JSON, no Markdown fences and no commentary.

Important source hierarchy:
1. SOURCE CV TEXT is the authoritative source and must drive the JSON.
2. OFFICIAL CV XML is only supplemental for recovering formal titles, dates, publications, IDs, and metadata that are also consistent with the CV.
3. PROFILE BRIEF is only a short context note for tone and organization. Do not use it to replace the CV or invent CV entries.

Extraction rules:
- Do not produce a brief profile summary when detailed entries exist.
- Extract all visible education items, jobs, roles, projects, grants, awards, teaching, software, certifications, languages, identifiers, and technical skills that are present in the SOURCE CV or official XML.
- Publication, supervision, and event-heavy Ciência Vitae sections may be completed by the local PDF parser after this LLM step; focus here on accurate core profile extraction.
- Keep entries granular. Do not collapse many publications into "23 publications" if titles are available.
- Translate labels and descriptions to English, but keep proper nouns, journal names, project names, institutions, identifiers, and URLs intact.
- Prefer arrays of bullet strings for rich descriptions when useful.
- If a date range, DOI, URL, author role, funding code, or classification is visible, preserve it.

Required top-level schema. You may add extra fields inside objects when the source contains them:
{{
  "personal_info": {{}},
  "education": [{{"degree": "", "institution": "", "year": "", "description": ""}}],
  "experience": [{{"title": "", "company": "", "years": "", "description": ""}}],
  "publications": [{{"title": "", "journal": "", "year": ""}}],
  "projects": [{{"name": "", "role": "", "description": ""}}],
  "skills": ["skill_1", "skill_2"],
  "awards": [],
  "teaching": [],
  "supervision": [],
  "software": [],
  "languages": [],
  "certifications": [],
  "events": []
}}

Professional brief for tone only:
Dany Dominguez Perez, PhD. Hybrid profile across bioinformatics/omics, biodiscovery, scientific B2B account management, and solution-oriented scientific-commercial writing.

SOURCE CV TEXT:
{cv_payload}

OFFICIAL CV XML SUPPLEMENT:
{xml_text}

PROFILE BRIEF:
{dossier}
"""
    try:
        parsed = normalize_cv_profile(llm_json(prompt, temperature=0.1))
    except Exception as exc:
        parsed = normalize_cv_profile(heuristic_cv_json(cv_text, dossier))
        parsed["source_notes"]["llm_error"] = str(exc)
    parsed = enrich_cv_from_source(parsed, cv_text)
    parsed = normalize_cv_profile(parsed)
    CV_JSON_PATH.write_text(json.dumps(parsed, ensure_ascii=False, indent=2), encoding="utf-8")
    return parsed


def language_hint(job_text: str) -> str:
    text = job_text.lower()
    pt_hits = len(re.findall(r"\b(vaga|candidato|experiencia|licenciatura|empresa|portugal|conhecimento)\b", text))
    es_hits = len(re.findall(r"\b(oferta|puesto|candidato|experiencia|empresa|licenciatura|conocimiento)\b", text))
    if pt_hits > es_hits and pt_hits >= 2:
        return "PT"
    if es_hits >= 2:
        return "ES"
    return "EN"


def keyword_hints(job_text: str, limit: int = 18) -> list[str]:
    stop = {
        "and", "the", "for", "with", "you", "are", "our", "that", "this", "from", "will", "your", "have",
        "job", "role", "team", "work", "to", "of", "in", "a", "an", "or", "as", "is", "be", "we", "on",
        "de", "la", "el", "en", "para", "con", "una", "uno", "por", "que", "del", "los", "las", "um",
        "uma", "com", "das", "dos", "na", "no", "para",
    }
    tokens = re.findall(r"[A-Za-zÀ-ÿ][A-Za-zÀ-ÿ0-9+.#-]{2,}", job_text.lower())
    counts: dict[str, int] = {}
    for token in tokens:
        if token in stop or len(token) < 4:
            continue
        counts[token] = counts.get(token, 0) + 1
    return [word for word, _ in sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))[:limit]]


def heuristic_match(req: MatchRequest) -> dict[str, Any]:
    lang = language_hint(req.job_text)
    keywords = keyword_hints(req.job_text)
    cv = req.selected_cv or {}
    exp = cv.get("experience", [])
    skills = cv.get("skills", [])
    projects = cv.get("projects", [])
    pubs = cv.get("publications", [])
    skill_text = ", ".join(skills[:16]) if isinstance(skills, list) else str(skills)
    bullets = []
    for item in exp[:5]:
        if isinstance(item, dict):
            bullets.append(
                f"- {item.get('title', 'Relevant role')} at {item.get('company', 'relevant organization')}: "
                f"{item.get('description', '')} Keywords aligned: {', '.join(keywords[:5])}."
            )
    for item in projects[:3]:
        if isinstance(item, dict):
            bullets.append(
                f"- Project {item.get('name', 'relevant project')}: {item.get('description', '')}"
            )
    # Build a company/role-aware opening
    company_str = req.company or "your organization"
    position_str = req.position or "this position"
    keyword_str = ", ".join(keywords[:6]) if keywords else skill_text

    cover_templates = {
        "ES": f"""Estimado equipo de {company_str},

Me dirijo a ustedes con un genuino entusiasmo por la posición de {position_str}. La combinación de {keyword_str} que se busca encaja directamente con mi trayectoria y, sobre todo, con el tipo de impacto que quiero generar en esta etapa de mi carrera.

Soy doctor en Biología (Universidad de Porto, 2017) con más de una década de experiencia en bioinformática aplicada, proteómica, transcriptómica y descubrimiento de péptidos bioactivos. Soy autor principal de DeTox (Briefings in Bioinformatics, 2024), autor de FiltDeTox (EuVen 2023) y de SeqLengthPlot v2.0 (Bioinformatics Advances, 2024), y cuento con 43 publicaciones indexadas y 413 citas. Más allá de la ciencia básica, tengo experiencia comercial directa en ventas B2B de equipamiento científico (Proquinorte, Portugal), lo que me permite comunicar valor técnico a audiencias tanto investigadoras como industriales.

Dispongo de infraestructura HPC propia (160 hilos, 384 GB RAM), por lo que puedo abordar proyectos computacionalmente exigentes de forma inmediata y sin costes de nube adicionales para el cliente.

Me gustaría enormemente comentar cómo mi perfil puede contribuir a los objetivos de {company_str}. Quedo a su disposición para una llamada o entrevista cuando les sea conveniente.

Atentamente,
Dany Domínguez Pérez, PhD
PagBiOMICs · Porto, Portugal · discover@pagbiomics.com
ORCID: 0000-0002-5211-972X""",

        "PT": f"""Exma. equipa de recrutamento de {company_str},

Escrevo-vos com genuíno entusiasmo pela posição de {position_str}. O perfil que procuram — centrado em {keyword_str} — corresponde diretamente à minha trajetória e ao tipo de impacto que pretendo ter nesta fase da minha carreira.

Sou doutor em Biologia (Universidade do Porto, 2017) com mais de uma década de experiência em bioinformática aplicada, proteómica, transcriptómica e descoberta de péptidos bioativos. Sou autor principal do DeTox (Briefings in Bioinformatics, 2024), autor do FiltDeTox (EuVen 2023) e do SeqLengthPlot v2.0 (Bioinformatics Advances, 2024), com 43 publicações indexadas e 413 citações. Para além da investigação, tenho experiência comercial B2B em equipamento científico (Proquinorte, Portugal), o que me permite traduzir valor técnico para audiências tanto académicas como industriais.

Disponho de infraestrutura HPC própria (160 threads, 384 GB RAM) — o que significa que posso dar resposta a projetos computacionalmente exigentes de imediato.

Estaria muito grato pela oportunidade de conversar sobre como o meu perfil pode contribuir para os objetivos de {company_str}.

Com os melhores cumprimentos,
Dany Domínguez Pérez, PhD
PagBiOMICs · Porto, Portugal · discover@pagbiomics.com
ORCID: 0000-0002-5211-972X""",

        "EN": f"""Dear {company_str} Hiring Team,

I am writing with genuine excitement about the {position_str} opportunity. The focus on {keyword_str} maps directly onto work I have been doing for the past several years, and {company_str}'s mission is exactly the kind of environment where I want to contribute at this stage of my career.

My background: PhD in Biology (University of Porto, 2017), lead author of DeTox (Briefings in Bioinformatics, 2024), author of FiltDeTox (EuVen 2023), and lead author of SeqLengthPlot v2.0 (Bioinformatics Advances, 2024). I have 43 peer-reviewed publications and 413 citations, with work spanning proteotranscriptomics, NGS pipeline development, phylogenetics, and bioactive peptide discovery. My most recent role was Research Fellow at Stazione Zoologica Anton Dohrn, Naples (CRIMAC DEEPVEN project, 2024–2026), where I led proteomic characterization of deep-sea cnidarian venoms.

I run my own dual-Xeon HPC (160 threads, 384 GB RAM), meaning I can handle computationally heavy analyses immediately — no cloud provisioning delays. I work in English, Spanish (native), and Portuguese, which is a practical advantage for teams operating across Iberia, Latin America, or international consortia.

I would very much welcome the opportunity to speak with your team about how my profile fits what you are building.

Warm regards,
Dany Domínguez Pérez, PhD
PagBiOMICs | Porto, Portugal | discover@pagbiomics.com
ORCID: 0000-0002-5211-972X""",
    }
    cover = cover_templates[lang]
    questions = [
        {
            "question": "How would you prioritize between deep scientific validation and business delivery timelines?",
            "answer": "Frame the answer around risk-based milestones: define minimum evidence, communicate assumptions, and protect quality without slowing the whole decision cycle.",
        },
        {
            "question": "Your profile spans research and sales. Which identity is strongest for this role?",
            "answer": "Position the hybrid profile as a strength: scientific credibility plus customer-facing execution, adapted to the role type requested.",
        },
        {
            "question": "Which technical gap would you close first after joining?",
            "answer": "Name a specific tool or domain from the job description and explain a fast learning plan tied to a deliverable in the first 30 days.",
        },
    ]
    return {
        "adapted_cv": "\n".join(bullets) or "- Hybrid scientific and commercial profile aligned with the role requirements.",
        "cover_letter": cover,
        "interview_questions": questions,
        "keywords": keywords,
        "ats_score": min(92, 55 + len(set(keywords) & set(str(cv).lower().split())) * 4),
        "language": lang,
        "mode": "heuristic_fallback",
    }


def generate_match(req: MatchRequest) -> dict[str, Any]:
    refs = load_reference_bundle()
    company_str = req.company or "the company"
    position_str = req.position or "this role"
    prompt = f"""
You are a senior career strategist and ATS optimization expert writing on behalf of Dany Domínguez Pérez, PhD.

Generate a tailored application package by matching the selected CV data against the job offer.
Return only valid JSON with this schema:
{{
  "adapted_cv": "Markdown bullets. Use the offer keywords naturally. Keep claims factual and specific.",
  "cover_letter": "Persuasive letter in the language of the job offer (ES, PT, or EN). See tone rules below.",
  "interview_questions": [
    {{"question": "", "answer": ""}},
    {{"question": "", "answer": ""}},
    {{"question": "", "answer": ""}}
  ],
  "keywords": ["keyword"],
  "ats_score": 0,
  "language": "ES|PT|EN"
}}

=== CANDIDATE FACTS (use only these — never invent) ===
- PhD in Biology, University of Porto, 2017
- Research Fellow at Stazione Zoologica Anton Dohrn, Naples (CRIMAC DEEPVEN project, Feb 2024 – Jan 2026, CONCLUDED)
- Currently: Independent Bioinformatician / Freelancer (PagBiOMICs, Porto, since Feb 2026)
- Lead author: DeTox (Briefings in Bioinformatics, 2024); SeqLengthPlot v2.0 (Bioinformatics Advances, 2024)
- Author of FiltDeTox: open-source toxin candidate filtering tool (EuVen 2023, github.com/danydguezperez/FiltDeTox)
- 43 peer-reviewed publications; 413 citations; ORCID 0000-0002-5211-972X
- HPC: dual-node Xeon, 160 threads, 384 GB RAM, 9.3 TB storage (Porto)
- Languages: Spanish (native), Portuguese (fluent), English (professional)
- B2B sales experience: Proquinorte (scientific equipment, 2022–2023)
- Contact: discover@pagbiomics.com

=== TONE RULES FOR COVER LETTER ===
These rules are derived from the candidate's actual successful applications. Follow them strictly.

1. WARM AND SPECIFIC: Address the company/team by name in the opening. Show you know what they do.
   Bad: "Dear hiring team, I am applying for this opportunity..."
   Good: "Dear [Company] team, I am writing with genuine excitement about the [position] role. [Company]'s work on [specific thing from job description] is exactly the kind of environment..."

2. CONNECT SCIENCE TO COMMERCIAL VALUE: Explain how specific technical skills translate to business outcomes for this company. Don't just list skills — explain the impact.

3. MENTION THE HPC: Always include the private HPC infrastructure as a competitive advantage — no cloud delays, handles large datasets immediately.

4. CITE A SPECIFIC PUBLICATION when relevant: DeTox for pipeline/toxin/computational roles; ACS Omega AMP paper for drug discovery/peptide roles; SeqLengthPlot for visualization/NGS roles.

5. CLOSE WITH FORWARD MOMENTUM: End with a specific invitation to talk, not a generic "thank you for considering."

6. LANGUAGE: Write in the language of the job offer. For Spanish: use "usted" formality but warm tone. For Portuguese: European Portuguese.

7. SIGNATURE: Always end with full name, PhD, PagBiOMICs, Porto, Portugal, discover@pagbiomics.com, ORCID.

=== ROLE POSITIONING ===
Role type: {req.role_type}
- If Scientific: lead with DeTox, SeqLengthPlot, HPC, publications, pipeline engineering, non-model organism expertise
- If Commercial B2B: lead with Proquinorte experience, scientific credibility as sales differentiator, Iberian market knowledge, trilingual advantage
- If Hybrid: open with scientific credibility, pivot to "this is how I connect it to [company]'s commercial goals"

Position: {position_str}
Company: {company_str}
Job URL: {req.job_url or "Not provided"}

Selected CV JSON:
{json.dumps(req.selected_cv, ensure_ascii=False)}

Strategic profile dossier:
{refs["dossier"][:16000]}

Job offer text:
{req.job_text}
"""
    try:
        result = llm_json(prompt, temperature=0.35)
    except Exception as exc:
        result = heuristic_match(req)
        result["llm_error"] = str(exc)
    return result


def markdownish_to_doc(document: Any, title: str, content: str) -> None:
    document.add_heading(title, level=1)
    for raw in str(content).splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)


def application_markdown(
    adapted_cv: str,
    cover_letter: str,
    interview_questions: list[Any] | str,
    position: Optional[str] = None,
    company: Optional[str] = None,
) -> str:
    questions = interview_questions
    if not isinstance(questions, str):
        question_lines = []
        for index, item in enumerate(questions or [], start=1):
            if isinstance(item, dict):
                question_lines.append(f"### {index}. {item.get('question', 'Interview question')}\n\n{item.get('answer', '')}")
            else:
                question_lines.append(f"### {index}. {item}")
        questions = "\n\n".join(question_lines)
    return f"""# Tailored Application Package

**Position:** {position or "Target role"}  
**Company:** {company or "Target company"}  
**Generated:** {datetime.now().isoformat(timespec="seconds")}

## Tailored CV

{adapted_cv or ""}

## Cover Letter

{cover_letter or ""}

## Interview Preparation

{questions or ""}
"""


def save_markdown_package(
    app_id: str,
    result: dict[str, Any],
    position: Optional[str] = None,
    company: Optional[str] = None,
) -> Path:
    ensure_dirs()
    filename = f"application_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{app_id[:8]}.md"
    path = EXPORTS_DIR / filename
    path.write_text(
        application_markdown(
            result.get("adapted_cv", ""),
            result.get("cover_letter", ""),
            result.get("interview_questions", []),
            position=position,
            company=company,
        ),
        encoding="utf-8",
    )
    return path


def export_docx(payload: ExportRequest) -> Path:
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx is not installed. Run: pip install -r requirements.txt")
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("Tailored Application Package", 0)
    if payload.position or payload.company:
        doc.add_paragraph(f"{payload.position or 'Target role'} | {payload.company or 'Target company'}")
    markdownish_to_doc(doc, "Tailored CV", payload.adapted_cv)
    markdownish_to_doc(doc, "Cover Letter", payload.cover_letter)
    questions = payload.interview_questions
    if not isinstance(questions, str):
        questions = json.dumps(questions, ensure_ascii=False, indent=2)
    markdownish_to_doc(doc, "Interview Preparation", questions)
    filename = f"job_application_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = EXPORTS_DIR / filename
    doc.save(path)
    return path


def export_pdf(payload: ExportRequest) -> Path:
    if SimpleDocTemplate is None:
        raise HTTPException(status_code=500, detail="reportlab is not installed. Run: pip install -r requirements.txt")
    filename = f"job_application_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    path = EXPORTS_DIR / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = [Paragraph("Tailored Application Package", styles["Title"]), Spacer(1, 12)]
    for title, content in (
        ("Tailored CV", payload.adapted_cv),
        ("Cover Letter", payload.cover_letter),
        ("Interview Preparation", payload.interview_questions),
    ):
        story.append(Paragraph(title, styles["Heading1"]))
        text = content if isinstance(content, str) else json.dumps(content, ensure_ascii=False, indent=2)
        for paragraph in str(text).splitlines():
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles["BodyText"]))
                story.append(Spacer(1, 6))
    doc.build(story)
    return path


def export_markdown(payload: ExportRequest) -> Path:
    filename = f"job_application_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    path = EXPORTS_DIR / filename
    path.write_text(
        application_markdown(
            payload.adapted_cv,
            payload.cover_letter,
            payload.interview_questions,
            position=payload.position,
            company=payload.company,
        ),
        encoding="utf-8",
    )
    return path


def safe_filename(value: str, fallback: str = "parsed_cv") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", value or "").strip("._")
    return cleaned[:80] or fallback


def cv_value_markdown(value: Any, indent: int = 0) -> str:
    prefix = "  " * indent
    if isinstance(value, dict):
        lines = []
        for key, child in value.items():
            if child in (None, "", [], {}):
                continue
            child_text = cv_value_markdown(child, indent + 1)
            if "\n" in child_text:
                lines.append(f"{prefix}- **{title_case(key)}:**\n{child_text}")
            else:
                lines.append(f"{prefix}- **{title_case(key)}:** {child_text.strip()}")
        return "\n".join(lines)
    if isinstance(value, list):
        lines = []
        for child in value:
            child_text = cv_value_markdown(child, indent + 1).strip()
            if child_text:
                lines.append(f"{prefix}- {child_text}")
        return "\n".join(lines)
    return f"{prefix}{str(value).strip()}"


def title_case(value: str) -> str:
    return str(value).replace("_", " ").strip().title()


def parsed_cv_markdown(cv: dict[str, Any]) -> str:
    personal = cv.get("personal_info") or {}
    name = personal.get("full_name") or personal.get("name") or "Parsed CV"
    lines = [
        f"# {name}",
        "",
        f"_Parsed CV export generated {datetime.now().isoformat(timespec='seconds')}. This file is produced from structured CV data, without Ciencia Vitae page headers._",
        "",
    ]
    for section, value in cv.items():
        if section == "source_notes" or value in (None, "", [], {}):
            continue
        lines.append(f"## {title_case(section)}")
        rendered = cv_value_markdown(value).strip()
        lines.append(rendered or "")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def export_cv_json(payload: CVExportRequest) -> Path:
    ensure_dirs()
    filename = f"{safe_filename(payload.filename_prefix)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    path = EXPORTS_DIR / filename
    path.write_text(json.dumps(payload.cv, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def export_cv_markdown(payload: CVExportRequest) -> Path:
    ensure_dirs()
    filename = f"{safe_filename(payload.filename_prefix)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    path = EXPORTS_DIR / filename
    path.write_text(parsed_cv_markdown(payload.cv), encoding="utf-8")
    return path


def export_cv_docx(payload: CVExportRequest) -> Path:
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx is not installed. Run: pip install -r requirements.txt")
    ensure_dirs()
    filename = f"{safe_filename(payload.filename_prefix)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = EXPORTS_DIR / filename
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for line in parsed_cv_markdown(payload.cv).splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], 0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(re.sub(r"\*\*(.*?)\*\*", r"\1", stripped[2:]), style="List Bullet")
        else:
            doc.add_paragraph(re.sub(r"\*\*(.*?)\*\*", r"\1", stripped))
    doc.save(path)
    return path


def export_cv_pdf(payload: CVExportRequest) -> Path:
    if SimpleDocTemplate is None:
        raise HTTPException(status_code=500, detail="reportlab is not installed. Run: pip install -r requirements.txt")
    ensure_dirs()
    filename = f"{safe_filename(payload.filename_prefix)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    path = EXPORTS_DIR / filename
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []
    for line in parsed_cv_markdown(payload.cv).splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(escape(stripped[2:]), styles["Title"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(escape(stripped[3:]), styles["Heading1"]))
        else:
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            story.append(Paragraph(escape(clean), styles["BodyText"]))
        story.append(Spacer(1, 5))
    doc.build(story)
    return path


@app.on_event("startup")
def startup() -> None:
    init_db()


@app.get("/")
def root() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/pagbiomics-embed")
def pagbiomics_embed_preview() -> FileResponse:
    bundled = BASE_DIR / "pagbiomics_embed.html"
    local = _EXE_DIR / "pagbiomics_embed.html"
    return FileResponse(bundled if bundled.exists() else local)


@app.get("/api/health")
def health() -> dict[str, Any]:
    settings = load_llm_settings()
    provider_name, provider = active_provider_config()
    return {
        "ok": True,
        "provider": provider_name,
        "provider_label": provider.get("label") or provider_name,
        "model": provider.get("model"),
        "llm_configured": bool(provider.get("api_key")) or provider_can_run_without_key(provider_name),
        "gemini_configured": bool(settings.get("providers", {}).get("gemini", {}).get("api_key")),
        "local_cv_pdf": bool(first_existing_path(LOCAL_CV_PDF_CANDIDATES)),
        "local_cv_text": bool(first_existing_path(LOCAL_CV_TEXT_CANDIDATES)),
        "local_cv_xml": LOCAL_CV_XML.exists(),
        "profile_dossier": PROFILE_DOSSIER.exists(),
    }


@app.get("/api/providers")
def get_providers() -> dict[str, Any]:
    return public_llm_settings()


@app.post("/api/providers")
def save_provider(payload: ProviderSaveRequest) -> dict[str, Any]:
    settings = load_llm_settings()
    provider_name = payload.provider.strip()
    if provider_name not in settings["providers"]:
        raise HTTPException(status_code=400, detail=f"Unsupported provider '{provider_name}'.")
    provider = settings["providers"][provider_name]
    if payload.model.strip():
        provider["model"] = payload.model.strip()
    if payload.base_url is not None and payload.base_url.strip():
        provider["base_url"] = payload.base_url.strip().rstrip("/")
    if payload.clear_api_key:
        provider["api_key"] = ""
    if payload.api_key is not None and payload.api_key.strip() and "..." not in payload.api_key:
        provider["api_key"] = payload.api_key.strip()
    if payload.active_provider in settings["providers"]:
        settings["active_provider"] = payload.active_provider
    else:
        settings["active_provider"] = provider_name
    save_llm_settings(settings)
    return public_llm_settings(settings)


@app.post("/api/providers/test")
def test_provider(payload: ProviderTestRequest) -> dict[str, Any]:
    provider_name = payload.provider
    name, provider = active_provider_config(provider_name)
    try:
        text = llm_generate(
            "Reply with exactly this JSON: {\"ok\": true, \"message\": \"provider connected\"}",
            temperature=0,
            provider_name=name,
        )
        return {
            "ok": True,
            "provider": name,
            "model": provider.get("model"),
            "sample": text[:300],
        }
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/api/fetch-job")
async def fetch_job_url(payload: dict) -> dict:
    url = payload.get("url", "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="URL required")
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
            response = await client.get(url, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()[:15000]
        title = soup.find("title")
        title_str = title.get_text(" ", strip=True) if title else ""
        meta_title = soup.find("meta", property="og:title") or soup.find("meta", attrs={"name": "title"})
        if meta_title and meta_title.get("content"):
            title_str = meta_title["content"].strip()
        position, company = extract_job_title_company(title_str)
        return {
            "job_text": text,
            "company": company,
            "position": position,
            "title": title_str,
            "url": str(response.url),
        }
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not fetch URL: {exc}") from exc


def extract_job_title_company(title: str) -> tuple[str, str]:
    cleaned = re.sub(r"\s+", " ", title or "").strip()
    cleaned = re.sub(r"\s*\|\s*LinkedIn.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*-\s*LinkedIn.*$", "", cleaned, flags=re.IGNORECASE)
    if not cleaned:
        return "", ""
    for separator in (" | ", " - ", " — ", " – ", " at "):
        if separator in cleaned:
            left, right = cleaned.split(separator, 1)
            if separator.strip() == "at":
                return left.strip(), right.strip()
            return left.strip(), right.strip()
    return cleaned, ""


@app.post("/api/upload-cv")
async def upload_cv(
    file: UploadFile | None = File(default=None),
    use_local: bool = Form(default=False),
    use_ai: bool = Form(default=True),
) -> dict[str, Any]:
    ensure_dirs()
    if file and file.filename:
        suffix = Path(file.filename).suffix.lower()
        if suffix == ".pdf":
            text = extract_pdf_text(file)
        elif suffix == ".xml":
            raw = await file.read()
            text = extract_xml_text_from_string(raw.decode("utf-8", errors="replace"))
        else:
            raw = await file.read()
            text = raw.decode("utf-8", errors="replace")
        source = file.filename
    else:
        text, source = load_local_cv_text()

    if not text.strip():
        raise HTTPException(status_code=400, detail="No CV text could be extracted. Upload a PDF/TXT or check local CV path.")

    if use_ai:
        structured = cv_to_structured_json(text)
        parse_mode = "ai"
    else:
        structured = normalize_cv_profile(heuristic_cv_json(text, load_reference_bundle()["dossier"]))
        structured = enrich_cv_from_source(structured, text)
        structured = normalize_cv_profile(structured)
        CV_JSON_PATH.write_text(json.dumps(structured, ensure_ascii=False, indent=2), encoding="utf-8")
        parse_mode = "local_heuristic"
    return {
        "source": source,
        "parse_mode": parse_mode,
        "text_chars": len(text),
        "text_pages": len(re.findall(r"--- Page \d+ ---", text)),
        "cv": structured,
    }


@app.get("/api/cv")
def get_saved_cv() -> dict[str, Any]:
    if CV_JSON_PATH.exists():
        return {"cv": normalize_cv_profile(json.loads(read_text_file(CV_JSON_PATH)))}
    text, _source = load_local_cv_text()
    if not text:
        raise HTTPException(status_code=404, detail="No saved CV JSON and local CV text was not found.")
    return {"cv": cv_to_structured_json(text)}


@app.post("/api/match")
def match_job(req: MatchRequest) -> dict[str, Any]:
    result = generate_match(req)
    app_id = str(uuid.uuid4())
    saved_markdown = save_markdown_package(app_id, result, position=req.position, company=req.company)
    with db() as conn:
        conn.execute(
            """
            INSERT INTO applications
            (id, created_at, company, position, role_type, job_url, job_text, result_json)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                app_id,
                datetime.now().isoformat(timespec="seconds"),
                req.company,
                req.position,
                req.role_type,
                req.job_url,
                req.job_text,
                json.dumps(result, ensure_ascii=False),
            ),
        )
        conn.commit()
    return {"id": app_id, "result": result, "saved_markdown": str(saved_markdown)}


@app.post("/api/export")
def export_application(payload: ExportRequest) -> FileResponse:
    ensure_dirs()
    fmt = payload.format.lower().strip()
    if fmt == "pdf":
        path = export_pdf(payload)
    elif fmt == "md":
        path = export_markdown(payload)
    else:
        path = export_docx(payload)
    media_type = (
        "application/pdf"
        if path.suffix == ".pdf"
        else "text/markdown; charset=utf-8"
        if path.suffix == ".md"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    return FileResponse(path, media_type=media_type, filename=path.name)


@app.post("/api/export-cv")
def export_parsed_cv(payload: CVExportRequest) -> FileResponse:
    if not payload.cv:
        raise HTTPException(status_code=400, detail="No parsed CV data received.")
    payload.cv = normalize_cv_profile(payload.cv)
    fmt = payload.format.lower().strip()
    if fmt == "pdf":
        path = export_cv_pdf(payload)
        media_type = "application/pdf"
    elif fmt == "docx":
        path = export_cv_docx(payload)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "json":
        path = export_cv_json(payload)
        media_type = "application/json"
    else:
        path = export_cv_markdown(payload)
        media_type = "text/markdown; charset=utf-8"
    return FileResponse(path, media_type=media_type, filename=path.name)


@app.get("/api/history")
def history() -> dict[str, Any]:
    init_db()
    with db() as conn:
        rows = conn.execute(
            """
            SELECT id, created_at, company, position, role_type, job_url, job_text, result_json
            FROM applications
            ORDER BY created_at DESC
            LIMIT 100
            """
        ).fetchall()
    items = []
    for row in rows:
        result = json.loads(row["result_json"])
        items.append(
            {
                "id": row["id"],
                "created_at": row["created_at"],
                "company": row["company"],
                "position": row["position"],
                "role_type": row["role_type"],
                "job_url": row["job_url"],
                "job_excerpt": textwrap.shorten(row["job_text"], width=220, placeholder="..."),
                "ats_score": result.get("ats_score"),
                "language": result.get("language"),
                "result": result,
            }
        )
    return {"items": items}


@app.get("/api/open-exports")
def open_exports() -> dict[str, Any]:
    ensure_dirs()
    try:
        if os.name == "nt":
            os.startfile(EXPORTS_DIR)  # type: ignore[attr-defined]
        else:
            raise RuntimeError("Open exports folder is only supported automatically on Windows.")
        return {"ok": True, "path": str(EXPORTS_DIR)}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not open exports folder: {exc}") from exc


app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("jobapp_ai_assistant:app", host="127.0.0.1", port=8080, reload=True)
